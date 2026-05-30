from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from app.models.schemas import ReferralNoteDocxRequest
from app.services.context import require_session
from app.services.patient_resolver import get_patient_or_404


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", value.strip())
    cleaned = cleaned.strip("._")
    return cleaned[:120] or "referral_note"


def _default_filename(patient_name: str) -> str:
    date_stamp = datetime.utcnow().strftime("%Y%m%d")
    return _safe_filename(f"referral_note_{patient_name}_{date_stamp}.docx")


def _normalize_filename(filename: str | None, patient_name: str) -> str:
    if not filename:
        return _default_filename(patient_name)
    safe_name = _safe_filename(filename)
    if not safe_name.lower().endswith(".docx"):
        safe_name = f"{safe_name}.docx"
    return safe_name


def _add_metadata(document: Document, label: str, value: str | None) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value or "Not available")


def _add_note_body(document: Document, final_note: str) -> None:
    for raw_line in final_note.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue

        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if line.endswith(":") and len(line) <= 90:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            continue

        document.add_paragraph(line)


def build_referral_note_docx(token: str, request: ReferralNoteDocxRequest) -> tuple[bytes, str]:
    session = require_session(token, request.active_facility_id)
    patient = get_patient_or_404(session.active_facility_id, request.patient_id)

    document = Document()
    document.core_properties.title = "Referral Note"
    document.core_properties.subject = f"Referral note for {patient.full_name}"
    document.core_properties.author = "HealthStack Copilot"

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    document.add_heading("Referral Note", level=1)
    _add_metadata(document, "Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    _add_metadata(document, "Facility", session.active_facility_name or session.active_facility_id)
    _add_metadata(document, "Patient", patient.full_name)
    _add_metadata(document, "MRN", patient.mrn)
    _add_metadata(document, "HS ID", patient.hs_id)
    _add_metadata(document, "Patient ID", patient.patient_id)

    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.LINE)

    _add_note_body(document, request.final_note)

    document.add_paragraph()
    disclaimer = document.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This document was generated from a doctor-reviewed copilot draft and should be signed by the responsible clinician before use."
    )
    disclaimer_run.italic = True

    buffer = BytesIO()
    document.save(buffer)
    filename = _normalize_filename(request.filename, patient.full_name)
    return buffer.getvalue(), filename
